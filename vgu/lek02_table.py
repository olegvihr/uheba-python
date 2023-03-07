from docx import Document

document = Document()

document.add_heading("9 марта", 0)

# get table data
items = (
    (1, '1978', 'qweert'),
    (2, '1950', 'asdfgh'),
    (3, '1953', 'zxcvbn'),
    (4, '1981', 'pjkdfd'),
)
# add table
table = document.add_table(1, 3)

# populate header row
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Номер'
heading_cells[1].text = 'Год'
heading_cells[2].text = 'Событие'

# aad a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]

table.style = 'Colorful Shading Accent 1'
# table.style = 'Table Grid'

document.save('table_style.docx')
