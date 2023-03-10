from docx import Document
import docx.enum.text


document = Document()

document.add_heading('Заголовок документа', 0)
# document.add_heading('Заголовок первого уровня', level=1)

paragraph = document.add_paragraph('Жерков тронул шпорами лошадь, которая раза три, горячась, перебила ногами, не зная, с какой начать, справилась и поскакала, обгоняя роту и догоняя коляску, тоже в такт песни.')
paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY

document.add_paragraph('элемент ненумерованого списка', style='List Bullet')
document.add_paragraph('элемент нумерованого списка', style='List Number')
document.add_paragraph('элемент нумерованого списка', style='List Number')
document.add_paragraph('элемент нумерованого списка', style='List Number')
document.add_paragraph('абзац текста в виде цитаты', style='Intense Quote')


# p = document.add_paragraph('Абзац обычного текста этого документа')
# p.add_run('часть жырным шрифтом, ').bold = True
# p.add_run(' а часть ')
# p.add_run('курсивом.').italic = True

document.save('paragraph.docx')
