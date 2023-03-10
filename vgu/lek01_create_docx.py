from docx import Document
from docx.shared import Inches
import docx.enum.text

from docx2pdf import convert

txtfile = []
with open('ada.txt', encoding='utf-8') as file:
    title = file.readline()
    for line in file:
        txtfile.append(line)


doc = Document()
doc.add_heading(title, 0)
doc.add_picture('pic.jpeg', width=Inches(6.0))

for p in txtfile:
    if len(p) > 3:
        if p[:2] == '- ':
            doc.add_paragraph(p.strip('\n')[2:], style='List Bullet')
        elif '\n' in p:
            paragraphs = doc.add_paragraph(p.split('\n'))
            paragraphs.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif len(p) > 50:
            doc.add_paragraph(p.strip('\n'), style='Intense Quote')
        else:
            doc.add_heading(p.strip('\n'), level=1)
try:
    doc.save('Body.docx')
except PermissionError:
    print('Закройте файл')